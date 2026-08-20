from pathlib import Path
from collections import Counter

from openpyxl import load_workbook
from docx import Document

from BatkaAI.actions.file_actions import find_file


# =========================================================
# ПОИСК ДОКУМЕНТА
# =========================================================


def find_document(filename):
    matches = find_file(
        filename
    )

    if not matches:
        return None

    supported = {
        ".xlsx",
        ".docx",
    }

    for file_path in matches:
        if (
            file_path.suffix.lower()
            in supported
        ):
            return file_path

    return None


# =========================================================
# EXCEL
# =========================================================


def read_excel_structure(file_path):
    workbook = load_workbook(
        file_path,
        data_only=False
    )

    result = {
        "type": "excel",
        "path": str(file_path),
        "filename": file_path.name,
        "sheet_count": len(
            workbook.sheetnames
        ),
        "sheets": [],
    }

    for sheet in workbook.worksheets:

        max_row = sheet.max_row
        max_column = sheet.max_column

        formulas = []
        empty_cells = 0
        numeric_cells = 0
        text_cells = 0
        text_numbers = []
        populated_cells = 0

        row_values = []

        for row in sheet.iter_rows():
            current_row = []

            for cell in row:
                value = cell.value

                current_row.append(
                    value
                )

                if value is None:
                    empty_cells += 1
                    continue

                populated_cells += 1

                if (
                    isinstance(value, str)
                    and value.startswith("=")
                ):
                    formulas.append(
                        {
                            "cell": cell.coordinate,
                            "formula": value,
                        }
                    )

                elif isinstance(
                    value,
                    (int, float)
                ):
                    numeric_cells += 1

                elif isinstance(
                    value,
                    str
                ):
                    text_cells += 1

                    clean = (
                        value
                        .replace(" ", "")
                        .replace(",", ".")
                    )

                    try:
                        float(clean)

                        text_numbers.append(
                            cell.coordinate
                        )

                    except ValueError:
                        pass

            row_values.append(
                current_row
            )

        # ---------------------------------------------
        # Заголовки
        # ---------------------------------------------

        headers = []

        if max_row >= 1:
            for cell in sheet[1]:
                headers.append(
                    cell.value
                )

        # ---------------------------------------------
        # Дубликаты строк
        # ---------------------------------------------

        duplicate_rows = []

        normalized_rows = []

        for row_number, values in enumerate(
            row_values[1:],
            start=2
        ):
            normalized = tuple(
                str(value).strip().lower()
                if value is not None
                else ""
                for value in values
            )

            if any(normalized):
                normalized_rows.append(
                    (
                        row_number,
                        normalized,
                    )
                )

        counts = Counter(
            row
            for _, row
            in normalized_rows
        )

        for row_number, row in normalized_rows:
            if counts[row] > 1:
                duplicate_rows.append(
                    row_number
                )

        # ---------------------------------------------
        # Столбцы
        # ---------------------------------------------

        columns = []

        for column_number in range(
            1,
            max_column + 1
        ):
            letter = sheet.cell(
                row=1,
                column=column_number
            ).column_letter

            width = (
                sheet.column_dimensions[
                    letter
                ].width
            )

            columns.append(
                {
                    "index": column_number,
                    "letter": letter,
                    "header": (
                        sheet.cell(
                            row=1,
                            column=column_number
                        ).value
                    ),
                    "width": width,
                }
            )

        # ---------------------------------------------
        # Excel Tables
        # ---------------------------------------------

        tables = []

        for table in sheet.tables.values():
            tables.append(
                {
                    "name": table.displayName,
                    "range": table.ref,
                }
            )

        sheet_info = {
            "name": sheet.title,

            "max_row": max_row,
            "max_column": max_column,

            "populated_cells": (
                populated_cells
            ),

            "empty_cells": empty_cells,

            "numeric_cells": (
                numeric_cells
            ),

            "text_cells": text_cells,

            "text_numbers": (
                text_numbers
            ),

            "headers": headers,

            "columns": columns,

            "formula_count": len(
                formulas
            ),

            "formulas": formulas[:50],

            "duplicate_rows": sorted(
                set(duplicate_rows)
            ),

            "duplicate_count": len(
                set(duplicate_rows)
            ),

            "merged_ranges": [
                str(cell_range)
                for cell_range
                in sheet.merged_cells.ranges
            ],

            "freeze_panes": (
                str(sheet.freeze_panes)
                if sheet.freeze_panes
                else None
            ),

            "auto_filter": (
                sheet.auto_filter.ref
                if sheet.auto_filter.ref
                else None
            ),

            "tables": tables,
        }

        result["sheets"].append(
            sheet_info
        )

    workbook.close()

    return result


# =========================================================
# WORD
# =========================================================


def read_word_structure(file_path):
    document = Document(
        file_path
    )

    paragraphs = []
    headings = []

    empty_paragraphs = 0

    for index, paragraph in enumerate(
        document.paragraphs,
        start=1
    ):
        text = paragraph.text.strip()

        if not text:
            empty_paragraphs += 1
            continue

        style_name = (
            paragraph.style.name
            if paragraph.style
            else ""
        )

        paragraph_info = {
            "index": index,
            "text": text[:500],
            "style": style_name,
        }

        paragraphs.append(
            paragraph_info
        )

        if (
            style_name
            and style_name.lower().startswith(
                "heading"
            )
        ):
            headings.append(
                paragraph_info
            )

    tables = []

    for table_index, table in enumerate(
        document.tables,
        start=1
    ):
        rows = []

        for row in table.rows:
            rows.append(
                [
                    cell.text.strip()
                    for cell in row.cells
                ]
            )

        tables.append(
            {
                "index": table_index,
                "rows": len(
                    table.rows
                ),
                "columns": (
                    len(table.columns)
                    if table.rows
                    else 0
                ),
                "preview": rows[:10],
            }
        )

    inline_shapes = len(
        document.inline_shapes
    )

    result = {
        "type": "word",
        "path": str(file_path),
        "filename": file_path.name,

        "paragraph_count": len(
            document.paragraphs
        ),

        "non_empty_paragraphs": len(
            paragraphs
        ),

        "empty_paragraphs": (
            empty_paragraphs
        ),

        "heading_count": len(
            headings
        ),

        "headings": headings,

        "table_count": len(
            tables
        ),

        "tables": tables,

        "image_count": inline_shapes,

        "paragraph_preview": (
            paragraphs[:30]
        ),
    }

    return result


# =========================================================
# UNIVERSAL READER
# =========================================================


def read_document(filename):
    file_path = find_document(
        filename
    )

    if not file_path:
        return None

    extension = (
        file_path.suffix.lower()
    )

    if extension == ".xlsx":
        return read_excel_structure(
            file_path
        )

    if extension == ".docx":
        return read_word_structure(
            file_path
        )

    return None