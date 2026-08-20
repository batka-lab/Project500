import os
import tempfile
from pathlib import Path

from PIL import Image

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from BatkaAI.actions.file_actions import find_file


def _ensure_docx_name(filename):
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    return filename


def _find_word_file(filename):
    matches = find_file(filename)

    if not matches:
        return None

    for file_path in matches:
        if file_path.suffix.lower() == ".docx":
            return file_path

    return None


def _alignment(value):
    value = str(value or "").lower()

    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    return alignments.get(
        value,
        WD_ALIGN_PARAGRAPH.LEFT
    )


def _apply_run_format(run, data):
    if data.get("bold") is not None:
        run.bold = bool(data.get("bold"))

    if data.get("italic") is not None:
        run.italic = bool(data.get("italic"))

    if data.get("underline") is not None:
        run.underline = bool(data.get("underline"))

    font_size = data.get("font_size")

    if font_size:
        run.font.size = Pt(
            float(font_size)
        )


def _prepare_image(image_path):
    """
    Приводит любое поддерживаемое Pillow изображение
    к обычному PNG, который стабильно понимает python-docx.
    """

    try:
        image = Image.open(image_path)

        print(
            f"Исходный формат изображения: "
            f"{image.format}"
        )

        # Убираем потенциально проблемные режимы
        if image.mode not in ("RGB", "RGBA"):
            image = image.convert("RGB")

        temp_dir = Path(
            tempfile.gettempdir()
        )

        temp_path = (
            temp_dir
            / "batka_ai_word_image.png"
        )

        image.save(
            temp_path,
            format="PNG"
        )

        image.close()

        print(
            f"Изображение подготовлено для Word: "
            f"{temp_path}"
        )

        return temp_path

    except Exception as e:
        print(
            f"Ошибка подготовки изображения: "
            f"{type(e).__name__}: {repr(e)}"
        )

        return None


def word_edit(filename, operation, data=None):
    data = data or {}

    try:

        # ==================================================
        # CREATE
        # ==================================================

        if operation == "create":
            filename = _ensure_docx_name(
                filename
            )

            desktop = Path.home() / "Desktop"

            file_path = (
                desktop
                / filename
            )

            document = Document()

            title = data.get(
                "title",
                ""
            )

            if title:
                paragraph = (
                    document.add_paragraph()
                )

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                run = paragraph.add_run(
                    title
                )

                run.bold = True
                run.font.size = Pt(18)

            document.save(
                file_path
            )

            print(
                f"Word-документ создан: "
                f"{file_path}"
            )

            return file_path

        # ==================================================
        # FIND EXISTING DOCUMENT
        # ==================================================

        file_path = _find_word_file(
            filename
        )

        if not file_path:
            print(
                f"Word-документ не найден: "
                f"{filename}"
            )

            return None

        # ==================================================
        # OPEN
        # ==================================================

        if operation == "open":
            os.startfile(
                file_path
            )

            print(
                f"Word-документ открыт: "
                f"{file_path}"
            )

            return file_path

        document = Document(
            file_path
        )

        # ==================================================
        # READ
        # ==================================================

        if operation == "read":
            print(
                f"Содержимое Word-документа "
                f"{file_path}:"
            )

            print("-" * 60)

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()

                if text:
                    print(text)

            for table_number, table in enumerate(
                document.tables,
                start=1
            ):
                print()
                print(
                    f"[ТАБЛИЦА {table_number}]"
                )

                for row in table.rows:
                    values = [
                        cell.text
                        for cell in row.cells
                    ]

                    print(
                        " | ".join(values)
                    )

            print("-" * 60)

            return file_path

        # ==================================================
        # PARAGRAPH
        # ==================================================

        elif operation == "add_paragraph":
            text = data.get(
                "text",
                ""
            )

            paragraph = (
                document.add_paragraph()
            )

            paragraph.alignment = _alignment(
                data.get("alignment")
            )

            run = paragraph.add_run(
                text
            )

            _apply_run_format(
                run,
                data
            )

        # ==================================================
        # HEADING
        # ==================================================

        elif operation == "add_heading":
            text = data.get(
                "text",
                ""
            )

            level = int(
                data.get(
                    "level",
                    1
                )
            )

            level = max(
                1,
                min(level, 9)
            )

            paragraph = document.add_heading(
                text,
                level=level
            )

            paragraph.alignment = _alignment(
                data.get(
                    "alignment",
                    "center"
                )
            )

            for run in paragraph.runs:
                run.bold = True

                font_size = data.get(
                    "font_size"
                )

                if font_size:
                    run.font.size = Pt(
                        float(font_size)
                    )

        # ==================================================
        # BULLET LIST
        # ==================================================

        elif operation == "add_bullet_list":
            items = data.get(
                "items",
                []
            )

            for item in items:
                paragraph = (
                    document.add_paragraph(
                        style="List Bullet"
                    )
                )

                paragraph.add_run(
                    str(item)
                )

        # ==================================================
        # NUMBERED LIST
        # ==================================================

        elif operation == "add_numbered_list":
            items = data.get(
                "items",
                []
            )

            for item in items:
                paragraph = (
                    document.add_paragraph(
                        style="List Number"
                    )
                )

                paragraph.add_run(
                    str(item)
                )

        # ==================================================
        # TABLE
        # ==================================================

        elif operation == "add_table":
            headers = data.get(
                "headers",
                []
            )

            rows = data.get(
                "rows",
                []
            )

            columns = len(
                headers
            )

            if columns == 0 and rows:
                columns = len(
                    rows[0]
                )

            if columns == 0:
                print(
                    "Не указаны данные таблицы."
                )

                return None

            table = document.add_table(
                rows=1 if headers else 0,
                cols=columns
            )

            table.style = "Table Grid"

            table.alignment = (
                WD_TABLE_ALIGNMENT.CENTER
            )

            if headers:
                cells = (
                    table.rows[0].cells
                )

                for index, header in enumerate(
                    headers
                ):
                    cells[index].text = str(
                        header
                    )

                    for paragraph in (
                        cells[index].paragraphs
                    ):
                        for run in paragraph.runs:
                            run.bold = True

            for row_data in rows:
                cells = (
                    table.add_row().cells
                )

                for index in range(
                    min(
                        len(row_data),
                        columns
                    )
                ):
                    cells[index].text = str(
                        row_data[index]
                    )

        # ==================================================
        # ADD TABLE ROW
        # ==================================================

        elif operation == "add_table_row":
            table_index = int(
                data.get(
                    "table_index",
                    0
                )
            )

            values = data.get(
                "values",
                []
            )

            if table_index >= len(
                document.tables
            ):
                print(
                    "Таблица не найдена."
                )

                return None

            table = document.tables[
                table_index
            ]

            cells = (
                table.add_row().cells
            )

            for index in range(
                min(
                    len(values),
                    len(cells)
                )
            ):
                cells[index].text = str(
                    values[index]
                )

        # ==================================================
        # SET TABLE CELL
        # ==================================================

        elif operation == "set_table_cell":
            table_index = int(
                data.get(
                    "table_index",
                    0
                )
            )

            row = int(
                data.get(
                    "row",
                    0
                )
            )

            column = int(
                data.get(
                    "column",
                    0
                )
            )

            text = str(
                data.get(
                    "text",
                    ""
                )
            )

            if table_index >= len(
                document.tables
            ):
                print(
                    "Таблица не найдена."
                )

                return None

            table = document.tables[
                table_index
            ]

            if row >= len(
                table.rows
            ):
                print(
                    "Строка таблицы не найдена."
                )

                return None

            if column >= len(
                table.columns
            ):
                print(
                    "Столбец таблицы не найден."
                )

                return None

            table.cell(
                row,
                column
            ).text = text

        # ==================================================
        # REPLACE TEXT
        # ==================================================

        elif operation == "replace_text":
            old_text = str(
                data.get(
                    "old_text",
                    ""
                )
            )

            new_text = str(
                data.get(
                    "new_text",
                    ""
                )
            )

            if not old_text:
                print(
                    "Не указан текст для замены."
                )

                return None

            replacements = 0

            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = (
                            run.text.replace(
                                old_text,
                                new_text
                            )
                        )

                        replacements += 1

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = (
                                        run.text.replace(
                                            old_text,
                                            new_text
                                        )
                                    )

                                    replacements += 1

            print(
                f"Выполнено замен: "
                f"{replacements}"
            )

        # ==================================================
        # FORMAT TEXT
        # ==================================================

        elif operation == "format_text":
            target = str(
                data.get(
                    "target",
                    ""
                )
            )

            found = 0

            for paragraph in document.paragraphs:
                if target.lower() in (
                    paragraph.text.lower()
                ):
                    paragraph.alignment = _alignment(
                        data.get(
                            "alignment"
                        )
                    )

                    for run in paragraph.runs:
                        _apply_run_format(
                            run,
                            data
                        )

                    found += 1

            print(
                f"Отформатировано абзацев: "
                f"{found}"
            )

        # ==================================================
        # IMAGE
        # ==================================================

        elif operation == "add_image":
            image_name = str(
                data.get(
                    "image",
                    ""
                )
            ).strip()

            if not image_name:
                print(
                    "Не указано имя изображения."
                )

                return None

            image_matches = find_file(
                image_name
            )

            if not image_matches:
                print(
                    f"Изображение не найдено: "
                    f"{image_name}"
                )

                return None

            image_path = image_matches[0]

            print(
                f"Найдено изображение: "
                f"{image_path}"
            )

            prepared_image = _prepare_image(
                image_path
            )

            if not prepared_image:
                return None

            paragraph = (
                document.add_paragraph()
            )

            paragraph.alignment = _alignment(
                data.get(
                    "alignment",
                    "center"
                )
            )

            run = paragraph.add_run()

            width = data.get(
                "width_inches"
            )

            if width:
                run.add_picture(
                    str(prepared_image),
                    width=Inches(
                        float(width)
                    )
                )
            else:
                run.add_picture(
                    str(prepared_image)
                )

            print(
                "Изображение добавлено "
                "в документ."
            )

        # ==================================================
        # PAGE BREAK
        # ==================================================

        elif operation == "add_page_break":
            document.add_page_break()

        # ==================================================
        # DEFAULT FONT
        # ==================================================

        elif operation == "set_default_font":
            font_name = data.get(
                "font_name",
                "Times New Roman"
            )

            font_size = float(
                data.get(
                    "font_size",
                    14
                )
            )

            style = document.styles[
                "Normal"
            ]

            style.font.name = (
                font_name
            )

            style.font.size = Pt(
                font_size
            )

        # ==================================================
        # SAVE AS
        # ==================================================

        elif operation == "save_as":
            new_filename = str(
                data.get(
                    "new_filename",
                    ""
                )
            ).strip()

            if not new_filename:
                print(
                    "Не указано новое имя файла."
                )

                return None

            new_filename = _ensure_docx_name(
                new_filename
            )

            new_path = (
                Path.home()
                / "Desktop"
                / new_filename
            )

            document.save(
                new_path
            )

            print(
                f"Документ сохранён как: "
                f"{new_path}"
            )

            return new_path

        else:
            print(
                f"Неизвестная операция Word: "
                f"{operation}"
            )

            return None

        # ==================================================
        # SAVE
        # ==================================================

        document.save(
            file_path
        )

        print(
            f"Word-документ изменён: "
            f"{file_path}"
        )

        return file_path

    except PermissionError as e:
        print(
            "Не удалось сохранить документ. "
            "Закрой его в Microsoft Word "
            "и повтори команду."
        )

        print(
            f"Техническая ошибка: "
            f"{type(e).__name__}: {repr(e)}"
        )

        return None

    except Exception as e:
        print(
            f"Ошибка Word Engine: "
            f"{type(e).__name__}: {repr(e)}"
        )

        return None


# ==================================================
# СТАРЫЕ WORD-КОМАНДЫ
# ==================================================


def create_word(filename, content):
    file_path = word_edit(
        filename,
        "create",
        {}
    )

    if file_path and content:
        word_edit(
            filename,
            "add_paragraph",
            {
                "text": content
            }
        )

    return file_path


def read_word(filename):
    return word_edit(
        filename,
        "read",
        {}
    )


def append_word(filename, content):
    return word_edit(
        filename,
        "add_paragraph",
        {
            "text": content
        }
    )


def add_word_heading(filename, content):
    return word_edit(
        filename,
        "add_heading",
        {
            "text": content,
            "level": 1,
            "alignment": "center",
            "font_size": 18
        }
    )